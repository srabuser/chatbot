from docx import Document

def read_courses_from_word(file_path):
    doc = Document(file_path)
    courses = []
    for table in doc.tables:
        for row in table.rows:
            course_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(course_data) == 4:  # Adjust based on your actual data structure
                course_code, course_name, credits, prerequisites = course_data
                courses.append({
                    'course_code': course_code,
                    'course_name': course_name,
                    'credits': credits,
                    'prerequisites': prerequisites
                })
    return courses
from docx import Document

def read_courses_from_word(file_path):
    doc = Document(file_path)
    courses = []
    for table in doc.tables:
        for row in table.rows:
            course_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(course_data) == 4:  # Adjust based on your actual data structure
                course_code, course_name, credits, prerequisites = course_data
                courses.append({
                    'course_code': course_code,
                    'course_name': course_name,
                    'credits': credits,
                    'prerequisites': prerequisites
                })
    return courses

def preprocess_course_input(course_input):
    # Normalize input (lowercase and remove extra spaces)
    course_input = course_input.strip().lower()
    # Convert Arabic numerals to English numerals (if applicable)
    arabic_to_english_map = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
    course_input = course_input.translate(arabic_to_english_map)
    return course_input

def find_course_prerequisites(course_code, courses):
    course_code = preprocess_course_input(course_code)
    for course in courses:
        if preprocess_course_input(course['course_code']) == course_code or preprocess_course_input(course['course_name']) == course_code:
            return course['prerequisites']
    return "Course not found. Please check the course code or name."

# Load courses from the Word file
file_path = "data.docx"  # Replace with the correct path to your Word file
courses = read_courses_from_word(file_path)

# Test: Check prerequisites for a course
course_code = "عال ١١٤"  # Example input
prerequisites = find_course_prerequisites(course_code, courses)
print(f"The prerequisites for {course_code} are: {prerequisites}")

# Test reading the courses
file_path = "data.docx"  # Update to the correct file path
courses = read_courses_from_word(file_path)
for course in courses:
    print(course)
